# !/usr/bin/env python3
# Title: DocumentScanner
###Description: Word counter for sentences in Costa Rican legal texts, aids lawyers identified complex structures that undermines reader comprehension.
##Author: Luis Ernesto Garcia Estrada
##Created: 5th September 2020
##Update: 11th November 2020

import io
import glob
import re
from docx import Document
from docx.shared import RGBColor


def oracionSplit(text):
    """function that receives a string from a read file and follows formatting steps: 1. remove unwanted line breaks, newlines, tabs, trailing spaces and replaces a common accronym in legal discourse 2. parses text on full stops and space 3. counts amount of words in every sentence split, avoids unwanted items that might include only symbols or pesky trailing white spaces. 4. returns a list"""
    word_count = 0
    line_count = 1

    text = re.sub("\t", "", text)
    text = re.sub("\n", " ", text)
    text = re.sub("\s+", " ", text)
    text = re.sub("\.\S", " ", text)
    text = re.sub("\s$", "", text)
    text = re.sub("Lic\.", "licenciado", text)
    text = re.sub("lic\.", "licenciado", text)

    parsed = text.split(". ")
    sentence_dict = {}

    for sentence in parsed:
        sentence = re.sub("\s+", " ", sentence)
        words = sentence.split(" ")
        word_count = len(words)
        if len(words) < 2:
            print("empty slot in line: " + str(line_count) + "; line skipped in document ")
        else:
            sentence_dict[sentence] = word_count

        line_count += 1
    return sentence_dict


def wordWriter(sentence_dict):
    document = Document()

    for sentence in sentence_dict.keys():
        # 4 a 30 bien
        if sentence_dict[sentence] < 30:
            p = document.add_paragraph()
            wp = p.add_run(sentence)
        # 30 a 60 mas o menos
        elif 30 <= sentence_dict[sentence] < 60:
            p = document.add_paragraph()
            wp = p.add_run(sentence)
            wp.font.color.rgb = RGBColor(0, 0, 255)
        # mas de 90 chao
        elif 60 <= sentence_dict[sentence] < 90:
            p = document.add_paragraph()
            wp = p.add_run(sentence)
            wp.font.color.rgb = RGBColor(0, 255, 0)
        else:
            p = document.add_paragraph()
            wp = p.add_run(sentence)
            wp.font.color.rgb = RGBColor(255, 0, 0)
    document.save('output.docx')
    pass


def main():
    filename = "xxxxxx"

    text = ""

    sentence_count = {}
    # try:
    with open(filename, "r") as f:
        try:
            text = f.read()
            sentence_list = oracionSplit(text)
            wordWriter(sentence_list)
        except UnicodeDecodeError:
            print("error in " + filename + " encoding please check it is utf-8")

    pass


if __name__ == '__main__':
    main()
    pass
